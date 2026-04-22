import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path
import json
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / 'data'
FIG = ROOT / 'figures'
DOCS = ROOT / 'docs'
DOCS.mkdir(parents=True, exist_ok=True)

SUMMARY = json.loads((DATA / 'analysis_summary.json').read_text(encoding='utf-8'))
PRE = pd.read_csv(DATA / 'preprocessing_summary.csv')
DESC = pd.read_csv(DATA / 'descriptive_statistics.csv', index_col=0)
RAW_SAMPLE = pd.read_csv(DATA / 'raw_sample_20rows.csv').head(8)
DICT = pd.read_csv(DATA / 'data_dictionary.csv')

BLUE = RGBColor(21, 76, 121)
DARK = RGBColor(32, 37, 43)
TEAL = RGBColor(58, 123, 154)
LIGHT = RGBColor(236, 243, 248)
GRAY = RGBColor(94, 103, 112)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=10.5, color=DARK, font='Noto Sans CJK KR'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styled_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE
    run.font.name = 'Noto Sans CJK KR'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)
        r2.font.color.rgb = GRAY
        r2.font.name = 'Noto Sans CJK KR'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')


def style_paragraph(paragraph, font_size=10.5, color=DARK, bold=False, space_after=6, line_spacing=1.35):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        run.bold = bold or run.bold
        run.font.name = 'Noto Sans CJK KR'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    style_paragraph(p)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
        style_paragraph(p, font_size=10.5, space_after=2)


def add_table_from_df(doc, df, title=None, column_widths=None, font_size=9.5):
    if title:
        cap = doc.add_paragraph()
        cap_run = cap.add_run(title)
        cap_run.bold = True
        cap_run.font.size = Pt(10.5)
        cap_run.font.color.rgb = BLUE
        cap_run.font.name = 'Noto Sans CJK KR'
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
        cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], col, bold=True, font_size=9.5, color=RGBColor(255, 255, 255))
        set_cell_shading(hdr[i], '154C79')
        if column_widths:
            hdr[i].width = Cm(column_widths[i])

    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, font_size=font_size)
            if column_widths:
                cells[i].width = Cm(column_widths[i])
        if len(table.rows) % 2 == 0:
            for c in cells:
                set_cell_shading(c, 'F4F8FB')
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, '1F2937')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    lines = code.strip().split('\n')
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        if idx < len(lines) - 1:
            run.add_break()
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()


def add_image(doc, path, width_cm, caption=None):
    doc.add_picture(str(path), width=Cm(width_cm))
    last_par = doc.paragraphs[-1]
    last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(cap, font_size=9, color=GRAY, space_after=8)


def set_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    normal = doc.styles['Normal']
    normal.font.name = 'Noto Sans CJK KR'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
    normal.font.size = Pt(10.5)

    for level, size in [(1, 15), (2, 12.5), (3, 11)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Noto Sans CJK KR'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE


def add_cover_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'EEF5FA')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run('중요 메모\n')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE
    r1.font.name = 'Noto Sans CJK KR'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
    r2 = p.add_run('현재 내부 DB에서 비영(非0) 실측 구간 확보가 어려워, 본 초안은 내부 스키마를 반영한 시뮬레이션 데이터로 분석 파이프라인을 검증한 버전이다. 실제 측정 데이터가 확보되면 동일 코드로 즉시 교체 가능하다.')
    r2.font.size = Pt(10.2)
    r2.font.color.rgb = DARK
    r2.font.name = 'Noto Sans CJK KR'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
    doc.add_paragraph()


def build_report():
    doc = Document()
    set_styles(doc)

    add_styled_title(doc, '레이더 기반 환자 생체신호 및 재실 상태 분석', '데이터마이닝 W08-W10 중간 발표 보고서 초안')
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = 'Table Grid'
    pairs = [
        ('작성자', '[이름 입력]'),
        ('디바이스 ID', '661'),
        ('분석 기간', '2025-03-01 ~ 2025-03-14 (14일, 1분 단위)'),
        ('데이터 형태', 'MySQL 스키마를 반영한 long format 시뮬레이션 데이터'),
    ]
    for row, (k, v) in zip(meta.rows, pairs):
        set_cell_text(row.cells[0], k, bold=True, font_size=10)
        set_cell_shading(row.cells[0], 'EAF1F7')
        set_cell_text(row.cells[1], v, font_size=10)
    doc.add_paragraph()
    add_cover_callout(doc)

    add_heading(doc, '1. 프로젝트 개요', level=1)
    add_body(doc, '본 프로젝트의 목적은 레이더 센서 기반 환자 모니터링 시스템에서 수집되는 심박수(Heart), 호흡수(Breath), 낙상 이벤트(Drop), 재실 상태(Status) 데이터를 이용해 환자 상태의 시간적 패턴을 분석하는 것이다. 특히 부재 상태와 재실 상태를 구분하고, 낙상 이벤트 전후의 생체신호 변화를 탐색하여 향후 위험 상태 분류 모델의 기반을 마련하는 데 초점을 두었다.')
    add_bullets(doc, [
        '연구 질문 1: 재실/부재 상태를 안정적으로 구분할 수 있는 데이터 전처리 규칙은 무엇인가?',
        '연구 질문 2: 심박수와 호흡수는 시간대 및 상태 변화에 따라 어떤 분포와 상관관계를 보이는가?',
        '연구 질문 3: 낙상 이벤트 전후에 관측 가능한 생체신호 변화 패턴이 존재하는가?',
    ])

    add_heading(doc, '2. 데이터 수집 결과', level=1)
    add_body(doc, '실제 내부 테이블은 timestamp, device_id, data_category, min_value, avg_value, max_value 형태의 long format으로 저장된다고 가정하였다. 이 구조를 그대로 반영하여 시뮬레이션 데이터를 생성했고, 이후 동일한 SQL과 Python 전처리 코드로 분석 가능한지 검증하였다.')
    add_table_from_df(doc, DICT.rename(columns={'name': 'field_name'}), '표 1. data_category 매핑', [2.5, 3.5, 9.5])
    collection_df = pd.DataFrame([
        ['총 원시 row 수', f"{SUMMARY['raw_rows']:,}"],
        ['timestamp 수', f"{SUMMARY['wide_rows_before_cleaning']:,}"],
        ['재실 minute 수', f"{SUMMARY['present_minutes_before_cleaning']:,}"],
        ['부재 minute 수', f"{SUMMARY['absent_minutes_before_cleaning']:,}"],
        ['낙상 이벤트 수', f"{SUMMARY['drop_events_after_cleaning']:,}"],
    ], columns=['지표', '값'])
    add_table_from_df(doc, collection_df, '표 2. 수집 데이터 규모', [6.0, 4.0])

    add_body(doc, '내부 DB에서 데이터를 조회할 때 사용할 수 있는 기본 SQL 예시는 아래와 같다. 실제 운영 테이블에서 동일 쿼리를 적용하면 long format 원시 데이터를 추출할 수 있다.')
    add_code_block(doc, '''SELECT
    timestamp,
    device_id,
    data_category,
    min_value,
    avg_value,
    max_value
FROM bullscare_db.acq_data_1e
WHERE device_id = 661
  AND data_category IN (14211, 14215, 14221, 14223)
ORDER BY timestamp;''')
    add_table_from_df(doc, RAW_SAMPLE, '표 3. 원시 데이터 샘플 (상위 8행)', [3.2, 1.8, 2.5, 2.2, 2.2, 2.2], font_size=8.6)

    add_heading(doc, '3. 데이터 전처리', level=1)
    add_body(doc, '원시 데이터는 data_category별로 행이 분리되어 있기 때문에, 분석을 위해 먼저 timestamp 기준 피벗(pivot)을 수행하여 Heart, Breath, Drop, Status를 한 행으로 통합했다. 이후 중복 row, 부분 누락 row, 부재 구간(Status=0), 센서 오류로 추정되는 이상치를 단계적으로 제거하였다.')
    add_bullets(doc, [
        '피벗 변환: long format → wide format',
        '결측 제거: Heart/Breath/Drop/Status 중 일부가 비어 있는 timestamp 삭제',
        '부재 제거: Status = 0 인 구간 제외',
        '이상치 제거: Heart 45~130 bpm, Breath 8~30 rpm 범위를 벗어나는 값 제거',
        '파생 변수 생성: 5분 이동평균, 직전 시점 대비 변화량(delta) 계산',
    ])
    add_table_from_df(doc, PRE.rename(columns={'step': '전처리 단계', 'rows': 'row 수', 'note': '설명'}), '표 4. 전처리 전후 비교', [5.2, 2.8, 7.0], font_size=8.8)
    add_image(doc, FIG / 'fig_01_preprocessing_funnel.png', 15.8, '그림 1. 전처리 단계별 데이터 축소')

    add_heading(doc, '4. 탐색적 데이터 분석(EDA)', level=1)
    add_body(doc, '정제된 데이터는 총 {}행이며, 심박수 평균은 {:.2f} bpm, 호흡수 평균은 {:.2f} rpm으로 나타났다. 심박수와 호흡수 사이의 상관계수는 {:.3f}로 계산되어, 활동량이 높아질수록 두 신호가 함께 증가하는 경향을 확인할 수 있었다.'.format(
        f"{SUMMARY['analysis_rows_after_cleaning']:,}", SUMMARY['heart_mean'], SUMMARY['breath_mean'], SUMMARY['heart_breath_correlation']))
    desc_df = DESC.reset_index().rename(columns={'index': '통계량', 'Heart': 'Heart', 'Breath': 'Breath'})
    add_table_from_df(doc, desc_df, '표 5. Heart/Breath 기초 통계량', [3.2, 3.0, 3.0], font_size=9.0)
    add_image(doc, FIG / 'fig_02_status_distribution.png', 9.5, '그림 2. 재실/부재 상태 분포')
    add_image(doc, FIG / 'fig_03_timeline_72h.png', 16.5, '그림 3. 72시간 구간에서의 심박/호흡 시계열')
    add_image(doc, FIG / 'fig_04_distributions.png', 16.0, '그림 4. Heart/Breath 분포 히스토그램')
    add_image(doc, FIG / 'fig_05_correlation_scatter.png', 9.3, '그림 5. Heart-Breath 상관관계 산점도')
    add_image(doc, FIG / 'fig_06_drop_event_window.png', 14.0, '그림 6. 낙상 이벤트 전후 20분 구간 변화')
    add_image(doc, FIG / 'fig_07_hourly_pattern.png', 14.5, '그림 7. 시간대별 평균 심박/호흡 패턴')

    add_heading(doc, '5. 중간 인사이트', level=1)
    add_bullets(doc, [
        '부재 구간은 Heart/Breath가 0으로 기록되므로, Status 필터링이 분석 품질에 결정적이다.',
        'Heart와 Breath는 강한 양의 상관관계(r=0.867)를 보여, 활동량 변화 또는 긴장 상태를 함께 반영할 가능성이 있다.',
        '낙상 이벤트 주변 구간에서 심박수와 호흡수의 단기 상승이 관측되어, 이벤트 전조 탐지(feature engineering)의 후보가 된다.',
        '시간대 평균을 보면 새벽에는 안정/수면 구간 특성이, 오전과 저녁에는 활동 구간 특성이 뚜렷하게 나타난다.',
    ])

    add_heading(doc, '6. 향후 계획', level=1)
    add_bullets(doc, [
        '실제 내부 DB에서 비영(非0) 측정 구간을 확보하여 동일 파이프라인에 바로 적용',
        'Heart/Breath 변화량, 이동평균, 낙상 전후 윈도우 특징량을 추가 생성',
        'Isolation Forest 또는 One-Class SVM 기반 이상 탐지 실험 수행',
        '정상/위험 구간 분류 모델과 간단한 대시보드 형태의 시각화로 확장',
    ])

    add_heading(doc, '7. 제출용 GitHub 리포지토리 구성안', level=1)
    add_body(doc, 'GitHub에는 data/, figures/, src/, docs/ 구조를 유지하고 README.md에 프로젝트 개요, 데이터 생성/전처리 방법, 실행 예시를 정리하면 된다. 본 산출물에는 업로드용 README 초안과 Python 코드, SQL 예시가 함께 포함되어 있다.')

    out_path = DOCS / 'datamining_midterm_report.docx'
    doc.save(out_path)
    print(f'[OK] saved {out_path}')

if __name__ == '__main__':
    build_report()
